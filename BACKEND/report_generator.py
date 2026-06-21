import io
import os
import pandas as pd
from datetime import datetime
from PIL import Image as PILImage

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, KeepTogether, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def clean_markdown_for_pdf(text):
    if not text:
        return ""
    cleaned = text
    cleaned = cleaned.replace("**", "<b>")
    parts = cleaned.split("<b>")
    new_parts = []
    for i, part in enumerate(parts):
        if i == 0:
            new_parts.append(part)
        elif i % 2 == 1:
            new_parts.append("<b>" + part)
        else:
            new_parts.append("</b>" + part)
    cleaned = "".join(new_parts)
    
    cleaned = cleaned.replace("*", "<i>")
    parts_i = cleaned.split("<i>")
    new_parts_i = []
    for i, part in enumerate(parts_i):
        if i == 0:
            new_parts_i.append(part)
        elif i % 2 == 1:
            new_parts_i.append("<i>" + part)
        else:
            new_parts_i.append("</i>" + part)
    cleaned = "".join(new_parts_i)

    cleaned = cleaned.replace("\r\n", "<br/>").replace("\n", "<br/>")
    return cleaned

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generar_interpretaciones_negocio(df_filtrado, meta_horas, ciclo_stats, sla_cumplimiento, lista_variantes, matriz_traspaso, slowest_cases):
    interpretaciones = {}
    
    interpretaciones['mapa_procesos'] = (
        "El mapa de procesos real (Grafo DFG) revela la secuencia física e histórica de los pedidos. "
        "Las conexiones más gruesas o con números elevados muestran el camino estándar de la operación. "
        "Las rutas alternativas reflejan desvíos de calidad y tiempos de espera en tránsito, que incrementan "
        "directamente los costos del proceso."
    )
    
    col_act = 'Actividad' if 'Actividad' in df_filtrado.columns else 'actividad'
    if col_act in df_filtrado.columns and not df_filtrado.empty:
        act_counts = df_filtrado[col_act].value_counts()
        if len(act_counts) > 0:
            top_act = act_counts.index[0]
            top_val = act_counts.values[0]
            sec_act_txt = ""
            if len(act_counts) > 1:
                sec_act_txt = f", seguida de la actividad <b>{act_counts.index[1]}</b> con {act_counts.values[1]} repeticiones"
            
            interpretaciones['frecuencia_actividades'] = (
                f"La actividad con mayor volumen de transacciones es <b>{top_act}</b> con <b>{top_val}</b> repeticiones{sec_act_txt}. "
                f"Esto indica que el esfuerzo operativo del equipo se concentra principalmente en estas fases iniciales o de control."
            )
        else:
            interpretaciones['frecuencia_actividades'] = "No hay datos suficientes para analizar la frecuencia de actividades."
    else:
        interpretaciones['frecuencia_actividades'] = "Columna de actividades no encontrada."
        
    col_emp = 'Empleado' if 'Empleado' in df_filtrado.columns else 'empleado'
    if col_emp in df_filtrado.columns and not df_filtrado.empty:
        emp_counts = df_filtrado[col_emp].value_counts()
        if len(emp_counts) > 0:
            top_emp = emp_counts.index[0]
            top_val = emp_counts.values[0]
            sec_emp_txt = ""
            if len(emp_counts) > 1:
                sec_emp_txt = f" y <b>{emp_counts.index[1]}</b> con {emp_counts.values[1]} registros"
            
            interpretaciones['carga_empleados'] = (
                f"Se identifica una distribución donde <b>{top_emp}</b> registra el mayor volumen de trabajo con <b>{top_val}</b> eventos, "
                f"seguido de cerca por{sec_emp_txt}. Se recomienda monitorear este balanceo para evitar sobrecargas de tareas."
            )
        else:
            interpretaciones['carga_empleados'] = "No hay datos suficientes para analizar la carga por operario."
    else:
        interpretaciones['carga_empleados'] = "Columna de operarios no encontrada."
        
    col_ts = 'Timestamp' if 'Timestamp' in df_filtrado.columns else 'fecha_hora'
    if col_ts in df_filtrado.columns and not df_filtrado.empty:
        df_ts = df_filtrado.copy()
        df_ts[col_ts] = pd.to_datetime(df_ts[col_ts])
        df_ts['date_str'] = df_ts[col_ts].dt.strftime('%Y-%m-%d')
        time_counts = df_ts.groupby('date_str').size().sort_index()
        if not time_counts.empty:
            peak_date = time_counts.idxmax()
            peak_val = time_counts.max()
            avg_val = round(time_counts.mean(), 1)
            interpretaciones['evolucion_temporal'] = (
                f"La operación registra un promedio de <b>{avg_val}</b> transacciones por día. "
                f"El volumen máximo se alcanzó el día <b>{peak_date}</b> con un pico de <b>{peak_val}</b> eventos procesados. "
                f"Esta variabilidad estacional sugiere la necesidad de una planeación flexible en picos de demanda."
            )
        else:
            interpretaciones['evolucion_temporal'] = "No hay datos suficientes para calcular la evolución temporal."
    else:
        interpretaciones['evolucion_temporal'] = "Columna de tiempo no encontrada."
        
    if ciclo_stats and 'promedio' in ciclo_stats:
        retraso_perc = round(100.0 - sla_cumplimiento, 2)
        interpretaciones['tiempos_ciclo'] = (
            f"El tiempo de ciclo promedio de procesamiento es de <b>{ciclo_stats.get('promedio', 0):.2f} horas</b>. "
            f"El nivel de cumplimiento del SLA operativo es del <b>{sla_cumplimiento:.2f}%</b>, lo que indica que "
            f"el <b>{retraso_perc}%</b> de los lotes superaron el umbral permitido de <b>{meta_horas} horas</b>, "
            f"con una duración máxima registrada de <b>{ciclo_stats.get('maxima', 0):.2f} horas</b>."
        )
    else:
        interpretaciones['tiempos_ciclo'] = "No hay datos de tiempos para calcular el cumplimiento de SLA."
        
    if lista_variantes:
        num_vars = len(lista_variantes)
        var1 = lista_variantes[0]
        var1_id = var1.get('id')
        var1_cob = var1.get('porcentaje') or var1.get('cobertura') or 0.0
        interpretaciones['variantes'] = (
            f"El proceso se ejecuta a través de <b>{num_vars}</b> variantes o secuencias de flujo diferentes. "
            f"La ruta principal (Variante <b>{var1_id}</b>) representa el <b>{var1_cob}%</b> de los casos. "
            f"La existencia de otras variantes secundarias revela desvíos o reprocesos que atentan contra la estandarización."
        )
    else:
        interpretaciones['variantes'] = "No hay datos de variantes de flujo disponibles."
        
    if matriz_traspaso and matriz_traspaso.get('values') and len(matriz_traspaso.get('values')) > 0:
        idx = matriz_traspaso.get('index', [])
        cols = matriz_traspaso.get('columns', [])
        vals = matriz_traspaso.get('values', [])
        
        max_trans = 0
        from_emp = ""
        to_emp = ""
        for i in range(len(idx)):
            for j in range(len(cols)):
                if vals[i][j] > max_trans and idx[i] != cols[j]:
                    max_trans = int(vals[i][j])
                    from_emp = idx[i]
                    to_emp = cols[j]
                    
        if max_trans > 0:
            interpretaciones['matriz_traspaso'] = (
                f"La matriz de interacciones de equipo (SNA) revela que el traspaso de tareas más recurrente "
                f"ocurre de <b>{from_emp}</b> hacia <b>{to_emp}</b> con <b>{max_trans}</b> entregas directas. "
                f"Este acoplamiento estrecho denota dependencias específicas entre estos dos operarios."
            )
        else:
            interpretaciones['matriz_traspaso'] = (
                "La matriz SNA muestra traspasos de trabajo dispersos o autotareas sin un acoplamiento crítico entre operarios."
            )
    else:
        interpretaciones['matriz_traspaso'] = "No hay registros de traspasos interactivos en el equipo."
        
    if slowest_cases and len(slowest_cases) > 0:
        peor_caso = slowest_cases[0]
        peor_id = peor_caso.get('id_caso')
        peor_dur = peor_caso.get('duracion')
        interpretaciones['lotes_criticos'] = (
            f"Se detalla el historial de lotes críticos. El caso <b>{peor_id}</b> registró la mayor "
            f"desviación con una duración extrema de <b>{peor_dur}</b>, representando una brecha significativa "
            f"respecto a la meta de <b>{meta_horas} horas</b>. Requiere auditoría de cuellos de botella específicos."
        )
    else:
        interpretaciones['lotes_criticos'] = "No se detectaron lotes con retraso crítico en el subconjunto de datos."
        
    return interpretaciones

def generar_graficos_matplotlib(df_filtrado, id_ejecucion, meta_horas):
    import uuid
    uid = uuid.uuid4().hex[:8]
    static_dir = os.path.join(os.path.dirname(__file__), "static")
    if not os.path.exists(static_dir):
        os.makedirs(static_dir)
        
    paths = {}
    
    try:
        col_act = 'Actividad' if 'Actividad' in df_filtrado.columns else 'actividad'
        act_counts = df_filtrado[col_act].value_counts()
        act_counts.index = act_counts.index.astype(str)
        act_counts = act_counts[act_counts > 0]
        
        fig, ax = plt.subplots(figsize=(8.0, 3.2))
        act_counts.plot(kind='bar', ax=ax, color='#8b5cf6', width=0.6)
        ax.set_title("Frecuencia de Actividades", fontsize=10, fontweight='bold', color='#1e3a8a')
        ax.tick_params(axis='x', rotation=30, labelsize=7)
        ax.tick_params(axis='y', labelsize=8)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_color('#cbd5e1')
        ax.spines['bottom'].set_color('#cbd5e1')
        ax.grid(axis='y', linestyle='--', alpha=0.3)
        plt.tight_layout()
        
        path_act = os.path.join(static_dir, f"chart_act_{id_ejecucion}_{uid}.png")
        fig.savefig(path_act, dpi=150)
        plt.close(fig)
        paths['actividades'] = path_act
    except Exception as e:
        print(f"Error en grafica de actividades: {e}")
        
    try:
        col_emp = 'Empleado' if 'Empleado' in df_filtrado.columns else 'empleado'
        emp_counts = df_filtrado[col_emp].value_counts()
        emp_counts.index = emp_counts.index.astype(str)
        emp_counts = emp_counts[emp_counts > 0].head(10)
        
        fig, ax = plt.subplots(figsize=(8.0, 3.2))
        emp_counts.plot(kind='bar', ax=ax, color='#ff7f0e', width=0.6)
        ax.set_title("Carga de Trabajo por Empleado", fontsize=10, fontweight='bold', color='#1e3a8a')
        ax.tick_params(axis='x', rotation=30, labelsize=7)
        ax.tick_params(axis='y', labelsize=8)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_color('#cbd5e1')
        ax.spines['bottom'].set_color('#cbd5e1')
        ax.grid(axis='y', linestyle='--', alpha=0.3)
        plt.tight_layout()
        
        path_emp = os.path.join(static_dir, f"chart_emp_{id_ejecucion}_{uid}.png")
        fig.savefig(path_emp, dpi=150)
        plt.close(fig)
        paths['empleados'] = path_emp
    except Exception as e:
        print(f"Error en grafica de empleados: {e}")
        
    try:
        col_ts = 'Timestamp' if 'Timestamp' in df_filtrado.columns else 'fecha_hora'
        df_ts = df_filtrado.copy()
        df_ts[col_ts] = pd.to_datetime(df_ts[col_ts])
        df_ts['date_str'] = df_ts[col_ts].dt.strftime('%Y-%m-%d')
        time_counts = df_ts.groupby('date_str').size().sort_index()
        
        fig, ax = plt.subplots(figsize=(9.5, 3.0))
        time_counts.plot(kind='line', ax=ax, color='#10b981', marker='o', linewidth=2, markersize=4)
        ax.fill_between(time_counts.index, time_counts.values, color='#10b981', alpha=0.1)
        ax.set_title("Volumen de Eventos en el Tiempo", fontsize=10, fontweight='bold', color='#1e3a8a')
        ax.tick_params(axis='x', rotation=45, labelsize=7)
        ax.tick_params(axis='y', labelsize=8)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_color('#cbd5e1')
        ax.spines['bottom'].set_color('#cbd5e1')
        ax.grid(axis='y', linestyle='--', alpha=0.3)
        plt.tight_layout()
        
        path_time = os.path.join(static_dir, f"chart_time_{id_ejecucion}_{uid}.png")
        fig.savefig(path_time, dpi=150)
        plt.close(fig)
        paths['timeline'] = path_time
    except Exception as e:
        print(f"Error en grafica timeline: {e}")
        
    try:
        from BACKEND import miner_api
        case_durations = miner_api.calcular_tiempos_ciclo(df_filtrado)
        if not case_durations.empty:
            fig, ax = plt.subplots(figsize=(8.0, 3.0))
            ax.hist(case_durations['Duracion_Horas'], bins=10, color='#d62728', edgecolor='white', rwidth=0.8)
            ax.set_title("Distribución de Duraciones (Histograma SLA)", fontsize=10, fontweight='bold', color='#1e3a8a')
            ax.tick_params(axis='x', labelsize=7)
            ax.tick_params(axis='y', labelsize=8)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['left'].set_color('#cbd5e1')
            ax.spines['bottom'].set_color('#cbd5e1')
            ax.grid(axis='y', linestyle='--', alpha=0.3)
            plt.tight_layout()
            
            path_sla = os.path.join(static_dir, f"chart_sla_{id_ejecucion}_{uid}.png")
            fig.savefig(path_sla, dpi=150)
            plt.close(fig)
            paths['sla_histogram'] = path_sla
    except Exception as e:
        print(f"Error en histograma SLA: {e}")
        
    return paths

def generar_reporte_pdf(metadata, kpis, alert, insights_text, ruta_grafo_frec, ruta_grafo_des, ciclo_stats, sla_cumplimiento, lista_variantes, matriz_traspaso, slowest_cases, chart_paths, interpretaciones=None):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    
    styles = getSampleStyleSheet()
    
    style_normal = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=13,
        textColor=colors.HexColor('#1f2937')
    )
    
    style_normal_bold = ParagraphStyle(
        'CustomNormalBold',
        parent=style_normal,
        fontName='Helvetica-Bold'
    )
    
    style_h1 = ParagraphStyle(
        'CustomH1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=15,
        textColor=colors.HexColor('#1e3a8a'),
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )

    style_title = ParagraphStyle(
        'Title',
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.white,
        alignment=1
    )

    style_subtitle = ParagraphStyle(
        'Subtitle',
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#93c5fd'),
        alignment=1
    )

    story = []
    
    header_data = [
        [Paragraph("📊 INFORME DE MINERÍA DE PROCESOS (VALERS)", style_title)],
        [Paragraph("Reporte Ejecutivo de Auditoría y Tiempos de Ciclo", style_subtitle)]
    ]
    header_table = Table(header_data, colWidths=[540])
    header_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#1e3a8a')),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 10),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 8))
    
    f_inc = metadata.get('f_inicio') or "Sin Filtro"
    f_fin = metadata.get('f_fin') or "Sin Filtro"
    var_sel = metadata.get('var_seleccionada') or "Todas"
    emps_count = len(metadata.get('emps_seleccionados')) if metadata.get('emps_seleccionados') else "Todos"
    
    meta_data = [
        [
            Paragraph(f"<b>Análisis:</b> {metadata.get('nombre_analisis')}", style_normal),
            Paragraph(f"<b>Fecha de Generación:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}", style_normal)
        ],
        [
            Paragraph(f"<b>Rango de Fechas:</b> {f_inc} a {f_fin}", style_normal),
            Paragraph(f"<b>Variante:</b> {var_sel} | <b>Operarios:</b> {emps_count}", style_normal)
        ]
    ]
    meta_table = Table(meta_data, colWidths=[270, 270])
    meta_table.setStyle(TableStyle([
        ('LINEBELOW', (0,-1), (-1,-1), 0.5, colors.HexColor('#d1d5db')),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('TOPPADDING', (0,0), (-1,-1), 4),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 8))
    
    costo_fmt = f"S/. {kpis.get('costo_total', 0.0):,.2f}"
    kpi_data = [
        [
            Paragraph("<b>📦 TOTAL LOTES</b><br/><font size=14 color='#1e3a8a'><b>" + str(kpis.get('total_lotes')) + "</b></font>", style_normal),
            Paragraph("<b>🔄 TOTAL EVENTOS</b><br/><font size=14 color='#1e3a8a'><b>" + str(kpis.get('total_eventos')) + "</b></font>", style_normal),
            Paragraph("<b>💰 COSTO TOTAL</b><br/><font size=14 color='#1e3a8a'><b>" + costo_fmt + "</b></font>", style_normal)
        ]
    ]
    kpi_table = Table(kpi_data, colWidths=[180, 180, 180])
    kpi_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f3f4f6')),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#e5e7eb')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e5e7eb')),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0), (-1,-1), 8),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(kpi_table)
    story.append(Spacer(1, 8))
    
    if alert.get('num_retrasados', 0) > 0:
        alert_bg = colors.HexColor('#fef2f2')
        alert_border = colors.HexColor('#fca5a5')
        alert_txt = f"<b>🚨 ALERTA DE SLA:</b> Se detectaron <b>{alert.get('num_retrasados')} lotes retrasados</b> (superan la meta de {alert.get('meta_horas')} hrs). El <b>{alert.get('porcentaje_critico')}%</b> de los procesos excede el umbral, con una demora promedio de <b>{alert.get('promedio_retraso')} horas</b>."
    else:
        alert_bg = colors.HexColor('#f0fdf4')
        alert_border = colors.HexColor('#86efac')
        alert_txt = f"<b>✅ SLA CUMPLIDO:</b> Todos los lotes procesados cumplen con el umbral establecido de <b>{alert.get('meta_horas')} horas</b>. No se detectan cuellos de botella críticos."
        
    alert_table = Table([[Paragraph(alert_txt, style_normal)]], colWidths=[540])
    alert_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), alert_bg),
        ('BOX', (0,0), (-1,-1), 1, alert_border),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('LEFTPADDING', (0,0), (-1,-1), 10),
        ('RIGHTPADDING', (0,0), (-1,-1), 10),
    ]))
    story.append(alert_table)
    story.append(Spacer(1, 8))
    
    if insights_text:
        insights_html = clean_markdown_for_pdf(insights_text)
        story.append(Paragraph("🤖 Análisis con Inteligencia Artificial (Gemini)", style_h1))
        
        insights_box = Table([[Paragraph(insights_html, style_normal)]], colWidths=[540])
        insights_box.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#eff6ff')),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#bfdbfe')),
            ('TOPPADDING', (0,0), (-1,-1), 8),
            ('BOTTOMPADDING', (0,0), (-1,-1), 8),
            ('LEFTPADDING', (0,0), (-1,-1), 10),
            ('RIGHTPADDING', (0,0), (-1,-1), 10),
        ]))
        story.append(insights_box)
        story.append(Spacer(1, 8))
        
    if ruta_grafo_frec and os.path.exists(ruta_grafo_frec):
        story.append(Paragraph("🗺️ Mapa de Procesos - Frecuencia (Grafo DFG)", style_h1))
        try:
            with PILImage.open(ruta_grafo_frec) as img:
                w, h = img.size
            max_w = 460
            ratio = max_w / float(w)
            final_h = int(h * ratio)
            
            if final_h > 160:
                final_h = 160
                ratio = final_h / float(h)
                max_w = int(w * ratio)
                
            graph_img = Image(ruta_grafo_frec, width=max_w, height=final_h)
            story.append(graph_img)
            
            if interpretaciones and 'mapa_procesos' in interpretaciones:
                story.append(Spacer(1, 4))
                story.append(Paragraph(f"<b>Interpretación del Mapa (Frecuencia):</b> {interpretaciones['mapa_procesos']}", style_normal))
        except Exception as e:
            story.append(Paragraph(f"<i>No se pudo renderizar la imagen del grafo de frecuencia: {e}</i>", style_normal))
        story.append(Spacer(1, 8))

    if ruta_grafo_des and os.path.exists(ruta_grafo_des):
        story.append(Paragraph("⏱️ Mapa de Procesos - Desempeño (Grafo DFG)", style_h1))
        try:
            with PILImage.open(ruta_grafo_des) as img:
                w, h = img.size
            max_w = 460
            ratio = max_w / float(w)
            final_h = int(h * ratio)
            
            if final_h > 160:
                final_h = 160
                ratio = final_h / float(h)
                max_w = int(w * ratio)
                
            graph_img = Image(ruta_grafo_des, width=max_w, height=final_h)
            story.append(graph_img)
            
            story.append(Spacer(1, 4))
            story.append(Paragraph("<b>Interpretación del Mapa (Desempeño):</b> El grafo de desempeño muestra la duración promedio acumulada entre las transiciones de actividades, lo que facilita la identificación visual de las demoras más críticas en el flujo de trabajo.", style_normal))
        except Exception as e:
            story.append(Paragraph(f"<i>No se pudo renderizar la imagen del grafo de desempeño: {e}</i>", style_normal))
        story.append(Spacer(1, 8))
            
    if 'actividades' in chart_paths:
        story.append(Paragraph("📊 Distribución y Carga Operativa - Frecuencia de Actividades", style_h1))
        act_img = Image(chart_paths['actividades'], width=440, height=176)
        story.append(act_img)
        if interpretaciones and 'frecuencia_actividades' in interpretaciones:
            story.append(Spacer(1, 4))
            story.append(Paragraph(f"<b>Interpretación:</b> {interpretaciones['frecuencia_actividades']}", style_normal))
        story.append(Spacer(1, 8))
        
    if 'empleados' in chart_paths:
        story.append(Paragraph("📊 Distribución y Carga Operativa - Carga de Trabajo por Empleado", style_h1))
        emp_img = Image(chart_paths['empleados'], width=440, height=176)
        story.append(emp_img)
        if interpretaciones and 'carga_empleados' in interpretaciones:
            story.append(Spacer(1, 4))
            story.append(Paragraph(f"<b>Interpretación:</b> {interpretaciones['carga_empleados']}", style_normal))
        story.append(Spacer(1, 8))
 
    if 'timeline' in chart_paths:
        story.append(Paragraph("📈 Evolución Temporal del Volumen de Trabajo", style_h1))
        time_img = Image(chart_paths['timeline'], width=460, height=145)
        story.append(time_img)
        if interpretaciones and 'evolucion_temporal' in interpretaciones:
            story.append(Spacer(1, 4))
            story.append(Paragraph(f"<b>Interpretación:</b> {interpretaciones['evolucion_temporal']}", style_normal))
        story.append(Spacer(1, 8))
        
    story.append(Paragraph("⏱️ Métricas de Tiempo de Ciclo y SLA", style_h1))
    stats_data = [
        [
            Paragraph("<b>Métrica</b>", style_normal_bold),
            Paragraph("<b>Valor (Horas)</b>", style_normal_bold)
        ],
        [Paragraph("Tiempo Promedio", style_normal), Paragraph(f"{ciclo_stats.get('promedio', 0):.2f} hrs", style_normal)],
        [Paragraph("Tiempo Mediana", style_normal), Paragraph(f"{ciclo_stats.get('mediana', 0):.2f} hrs", style_normal)],
        [Paragraph("Tiempo Mínimo", style_normal), Paragraph(f"{ciclo_stats.get('minima', 0):.2f} hrs", style_normal)],
        [Paragraph("Tiempo Máximo", style_normal), Paragraph(f"{ciclo_stats.get('maxima', 0):.2f} hrs", style_normal)],
        [Paragraph("Cumplimiento SLA", style_normal), Paragraph(f"{sla_cumplimiento:.2f}%", style_normal)]
    ]
    stats_table = Table(stats_data, colWidths=[140, 110])
    stats_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#e5e7eb')),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#d1d5db')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e5e7eb')),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f9fafb')]),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
    ]))
    
    stats_layout = Table([[stats_table]], colWidths=[270])
    stats_layout.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
        ('RIGHTPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(stats_layout)
    story.append(Spacer(1, 6))
    
    if interpretaciones and 'tiempos_ciclo' in interpretaciones:
        story.append(Paragraph(f"<b>Interpretación del SLA:</b> {interpretaciones['tiempos_ciclo']}", style_normal))
        story.append(Spacer(1, 8))
        
    if 'sla_histogram' in chart_paths:
        story.append(Paragraph("📊 Histograma de Distribución de Duraciones", style_h1))
        sla_hist_img = Image(chart_paths['sla_histogram'], width=440, height=165)
        story.append(sla_hist_img)
        story.append(Spacer(1, 8))

    if lista_variantes:
        story.append(Paragraph("🔄 Principales Variantes de Flujo Detectadas", style_h1))
        var_rows = [
            [
                Paragraph("<b>Var</b>", style_normal_bold),
                Paragraph("<b>Casos</b>", style_normal_bold),
                Paragraph("<b>Cob</b>", style_normal_bold),
                Paragraph("<b>Secuencia de Actividades</b>", style_normal_bold)
            ]
        ]
        
        for v in lista_variantes[:5]:
            act_text = v.get('actividades')
            if isinstance(act_text, list):
                act_text = " ➔ ".join(act_text)
            var_rows.append([
                Paragraph(f"<b>{v.get('id')}</b>", style_normal),
                Paragraph(str(v.get('cantidad') or v.get('casos')), style_normal),
                Paragraph(f"{v.get('porcentaje') or v.get('cobertura')}%", style_normal),
                Paragraph(act_text, style_normal)
            ])
            
        var_table = Table(var_rows, colWidths=[55, 45, 50, 390])
        var_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#e5e7eb')),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#d1d5db')),
            ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e5e7eb')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f9fafb')]),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ]))
        story.append(var_table)
        
        if interpretaciones and 'variantes' in interpretaciones:
            story.append(Spacer(1, 4))
            story.append(Paragraph(f"<b>Interpretación:</b> {interpretaciones['variantes']}", style_normal))
        story.append(Spacer(1, 8))

    if matriz_traspaso and matriz_traspaso.get('values'):
        story.append(Paragraph("🤝 Matriz de Traspaso (SNA)", style_h1))
        
        idx = matriz_traspaso.get('index', [])
        cols = matriz_traspaso.get('columns', [])
        vals = matriz_traspaso.get('values', [])
        
        original_len_cols = len(cols)
        original_len_idx = len(idx)
        if len(cols) > 8:
            cols = cols[:8]
        if len(idx) > 8:
            idx = idx[:8]
            
        grid_headers = [Paragraph("<b>De \\ A</b>", style_normal_bold)] + [Paragraph(f"<b>{c}</b>", style_normal_bold) for c in cols]
        grid_rows = [grid_headers]
        
        for i, row_name in enumerate(idx):
            row_cells = [Paragraph(f"<b>{row_name}</b>", style_normal)]
            for j in range(len(cols)):
                val = vals[i][j]
                row_cells.append(Paragraph(str(int(val)) if val > 0 else "-", style_normal))
            grid_rows.append(row_cells)
            
        num_cols = len(cols) + 1
        w_cell = min(540 / num_cols, 80)
        col_w = [85] + [w_cell] * len(cols)
        
        handover_table = Table(grid_rows, colWidths=col_w)
        handover_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#e5e7eb')),
            ('BACKGROUND', (0,0), (0,-1), colors.HexColor('#f3f4f6')),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#d1d5db')),
            ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e5e7eb')),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ALIGN', (1,0), (-1,-1), 'CENTER'),
            ('TOPPADDING', (0,0), (-1,-1), 2),
            ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ('LEFTPADDING', (0,0), (-1,-1), 2),
            ('RIGHTPADDING', (0,0), (-1,-1), 2),
        ]))
        story.append(KeepTogether([handover_table]))
        if original_len_cols > 8 or original_len_idx > 8:
            story.append(Paragraph("<font size=7 color='#6b7280'>* Nota: La matriz se ha truncado a las primeras 8 columnas y filas para su visualización en el reporte.</font>", style_normal))
            
        if interpretaciones and 'matriz_traspaso' in interpretaciones:
            story.append(Spacer(1, 4))
            story.append(Paragraph(f"<b>Interpretación:</b> {interpretaciones['matriz_traspaso']}", style_normal))
        story.append(Spacer(1, 8))

    if slowest_cases:
        story.append(KeepTogether([
            Paragraph("🚨 Detalle de Lotes Críticos (Mayor Duración)", style_h1),
            Spacer(1, 4)
        ]))
        
        slow_headers = [
            Paragraph("<b>ID Lote</b>", style_normal_bold),
            Paragraph("<b>Inicio</b>", style_normal_bold),
            Paragraph("<b>Fin</b>", style_normal_bold),
            Paragraph("<b>Duración</b>", style_normal_bold),
            Paragraph("<b>Estado SLA</b>", style_normal_bold)
        ]
        slow_rows = [slow_headers]
        for c in slowest_cases[:5]:
            status_style = style_normal
            if "retraso" in c.get('status').lower():
                status_style = ParagraphStyle('RedText', parent=style_normal, textColor=colors.HexColor('#dc2626'), fontName='Helvetica-Bold')
            slow_rows.append([
                Paragraph(c.get('id_caso'), style_normal),
                Paragraph(c.get('inicio'), style_normal),
                Paragraph(c.get('fin'), style_normal),
                Paragraph(c.get('duracion'), style_normal),
                Paragraph(c.get('status'), status_style)
            ])
            
        slow_table = Table(slow_rows, colWidths=[90, 110, 110, 80, 150])
        slow_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#e5e7eb')),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#d1d5db')),
            ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e5e7eb')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f9fafb')]),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ]))
        story.append(KeepTogether([slow_table]))
        
        if interpretaciones and 'lotes_criticos' in interpretaciones:
            story.append(Spacer(1, 4))
            story.append(Paragraph(f"<b>Interpretación:</b> {interpretaciones['lotes_criticos']}", style_normal))
        
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def add_interpretation_docx(doc, text):
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    
    parts = text.split("<b>")
    for i, part in enumerate(parts):
        if i == 0:
            run = p.add_run(part)
            run.font.size = Pt(9.5)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
        else:
            subparts = part.split("</b>")
            if len(subparts) > 0:
                run_bold = p.add_run(subparts[0])
                run_bold.font.size = Pt(9.5)
                run_bold.font.bold = True
                run_bold.font.italic = True
                run_bold.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            if len(subparts) > 1:
                run_normal = p.add_run("".join(subparts[1:]))
                run_normal.font.size = Pt(9.5)
                run_normal.font.italic = True
                run_normal.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

def set_section_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

def set_section_portrait(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

def generar_reporte_docx(metadata, kpis, alert, insights_text, ruta_grafo_frec, ruta_grafo_des, ciclo_stats, sla_cumplimiento, lista_variantes, matriz_traspaso, slowest_cases, chart_paths, interpretaciones=None):
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        set_section_portrait(section)
        
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    
    table_header = doc.add_table(rows=1, cols=1)
    table_header.allow_autofit = False
    table_header.columns[0].width = Inches(7.5)
    cell = table_header.rows[0].cells[0]
    cell.width = Inches(7.5)
    set_cell_background(cell, "1E3A8A")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("📊 INFORME DE MINERÍA DE PROCESOS (VALERS)")
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Reporte Ejecutivo de Auditoría y Tiempos de Ciclo")
    run2.font.italic = True
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x93, 0xC5, 0xFD)
    
    doc.add_paragraph()
    
    table_meta = doc.add_table(rows=2, cols=2)
    table_meta.allow_autofit = False
    table_meta.columns[0].width = Inches(3.75)
    table_meta.columns[1].width = Inches(3.75)
    for row in table_meta.rows:
        row.cells[0].width = Inches(3.75)
        row.cells[1].width = Inches(3.75)
    
    f_inc = metadata.get('f_inicio') or "Sin Filtro"
    f_fin = metadata.get('f_fin') or "Sin Filtro"
    var_sel = metadata.get('var_seleccionada') or "Todas"
    emps_count = len(metadata.get('emps_seleccionados')) if metadata.get('emps_seleccionados') else "Todos"
    
    m_runs = [
        (0, 0, "Análisis: ", metadata.get('nombre_analisis')),
        (0, 1, "Fecha de Generación: ", datetime.now().strftime('%d/%m/%Y %H:%M')),
        (1, 0, "Rango de Fechas: ", f"{f_inc} a {f_fin}"),
        (1, 1, f"Variante: {var_sel} | Operarios: ", str(emps_count))
    ]
    
    for r, c, label, value in m_runs:
        cell_m = table_meta.rows[r].cells[c]
        set_cell_margins(cell_m, top=60, bottom=60, left=100, right=100)
        p_m = cell_m.paragraphs[0]
        run_lbl = p_m.add_run(label)
        run_lbl.bold = True
        run_lbl.font.size = Pt(9.5)
        run_val = p_m.add_run(value)
        run_val.font.size = Pt(9.5)
        
    doc.add_paragraph()
    
    table_kpi = doc.add_table(rows=1, cols=3)
    table_kpi.allow_autofit = False
    for col in table_kpi.columns:
        col.width = Inches(2.5)
    for row in table_kpi.rows:
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(2.5)
        row.cells[2].width = Inches(2.5)
        
    costo_fmt = f"S/. {kpis.get('costo_total', 0.0):,.2f}"
    kpis_list = [
        ("📦 TOTAL LOTES", str(kpis.get('total_lotes'))),
        ("🔄 TOTAL EVENTOS", str(kpis.get('total_eventos'))),
        ("💰 COSTO TOTAL", costo_fmt)
    ]
    
    for i, (label, val) in enumerate(kpis_list):
        cell_k = table_kpi.rows[0].cells[i]
        set_cell_background(cell_k, "F3F4F6")
        set_cell_margins(cell_k, top=120, bottom=120, left=100, right=100)
        p_k = cell_k.paragraphs[0]
        p_k.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run_lbl = p_k.add_run(label + "\n")
        run_lbl.font.size = Pt(9)
        run_lbl.font.bold = True
        run_lbl.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
        
        run_v = p_k.add_run(val)
        run_v.font.size = Pt(14)
        run_v.font.bold = True
        run_v.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
    doc.add_paragraph()
    
    table_alert = doc.add_table(rows=1, cols=1)
    table_alert.allow_autofit = False
    table_alert.columns[0].width = Inches(7.5)
    cell_a = table_alert.rows[0].cells[0]
    cell_a.width = Inches(7.5)
    
    if alert.get('num_retrasados', 0) > 0:
        set_cell_background(cell_a, "FEF2F2")
        alert_txt = f"🚨 ALERTA DE SLA: Se detectaron {alert.get('num_retrasados')} lotes retrasados (superan la meta de {alert.get('meta_horas')} hrs). El {alert.get('porcentaje_critico')}% de los procesos excede el umbral, con una demora promedio de {alert.get('promedio_retraso')} horas."
    else:
        set_cell_background(cell_a, "F0FDF4")
        alert_txt = f"✅ SLA CUMPLIDO: Todos los lotes procesados cumplen con el umbral establecido de {alert.get('meta_horas')} horas. No se detectan cuellos de botella críticos."
        
    set_cell_margins(cell_a, top=100, bottom=100, left=150, right=150)
    p_a = cell_a.paragraphs[0]
    run_a = p_a.add_run(alert_txt)
    run_a.font.size = Pt(10)
    run_a.font.bold = True
    if alert.get('num_retrasados', 0) > 0:
        run_a.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
    else:
        run_a.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
        
    doc.add_paragraph()
    
    if insights_text:
        h = doc.add_paragraph()
        run_h = h.add_run("🤖 Análisis con Inteligencia Artificial (Gemini)")
        run_h.font.size = Pt(13)
        run_h.font.bold = True
        run_h.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        table_ins = doc.add_table(rows=1, cols=1)
        table_ins.allow_autofit = False
        table_ins.columns[0].width = Inches(7.5)
        cell_ins = table_ins.rows[0].cells[0]
        cell_ins.width = Inches(7.5)
        set_cell_background(cell_ins, "EFF6FF")
        set_cell_margins(cell_ins, top=120, bottom=120, left=150, right=150)
        
        p_ins = cell_ins.paragraphs[0]
        raw_text = insights_text.replace("\r\n", "\n")
        lines = raw_text.split("\n")
        for idx_l, line in enumerate(lines):
            if idx_l > 0:
                p_ins = cell_ins.add_paragraph()
            
            parts = line.split("**")
            for idx_p, part in enumerate(parts):
                run_p = p_ins.add_run(part)
                if idx_p % 2 == 1:
                    run_p.font.bold = True
                run_p.font.size = Pt(9.5)
                
    if ruta_grafo_frec and os.path.exists(ruta_grafo_frec):
        h_g = doc.add_paragraph()
        run_hg = h_g.add_run("🗺️ Mapa de Procesos - Frecuencia (Grafo DFG)")
        run_hg.font.size = Pt(13)
        run_hg.font.bold = True
        run_hg.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        p_desc = doc.add_paragraph()
        run_desc = p_desc.add_run("Este grafo representa la frecuencia de las rutas físicas e históricas del proceso de inventario.")
        run_desc.font.italic = True
        run_desc.font.size = Pt(9.5)
        
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p_img.add_run().add_picture(ruta_grafo_frec, width=Inches(3.8))
        except Exception as e:
            p_img.add_run(f"No se pudo insertar la imagen del grafo de frecuencia: {e}")
            
        if interpretaciones and 'mapa_procesos' in interpretaciones:
            add_interpretation_docx(doc, interpretaciones['mapa_procesos'])
            
    if ruta_grafo_des and os.path.exists(ruta_grafo_des):
        h_g2 = doc.add_paragraph()
        run_hg2 = h_g2.add_run("⏱️ Mapa de Procesos - Desempeño (Grafo DFG)")
        run_hg2.font.size = Pt(13)
        run_hg2.font.bold = True
        run_hg2.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        p_desc2 = doc.add_paragraph()
        run_desc2 = p_desc2.add_run("Este grafo muestra la duración promedio en horas de las transiciones entre actividades, resaltando las demoras.")
        run_desc2.font.italic = True
        run_desc2.font.size = Pt(9.5)
        
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p_img2.add_run().add_picture(ruta_grafo_des, width=Inches(3.8))
        except Exception as e:
            p_img2.add_run(f"No se pudo insertar la imagen del grafo de desempeño: {e}")
            
        add_interpretation_docx(doc, "El grafo de desempeño muestra la duración promedio acumulada entre las transiciones de actividades, lo que facilita la identificación visual de las demoras más críticas en el flujo de trabajo.")
            
    if 'actividades' in chart_paths:
        h_dist1 = doc.add_paragraph()
        run_hdist1 = h_dist1.add_run("📊 Distribución Operativa - Frecuencia de Actividades")
        run_hdist1.font.size = Pt(13)
        run_hdist1.font.bold = True
        run_hdist1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        p_act = doc.add_paragraph()
        p_act.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_act.add_run().add_picture(chart_paths['actividades'], width=Inches(6.0))
        
        if interpretaciones and 'frecuencia_actividades' in interpretaciones:
            add_interpretation_docx(doc, interpretaciones['frecuencia_actividades'])
            
    if 'empleados' in chart_paths:
        h_dist2 = doc.add_paragraph()
        run_hdist2 = h_dist2.add_run("📊 Distribución Operativa - Carga de Trabajo por Empleado")
        run_hdist2.font.size = Pt(13)
        run_hdist2.font.bold = True
        run_hdist2.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        p_emp = doc.add_paragraph()
        p_emp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_emp.add_run().add_picture(chart_paths['empleados'], width=Inches(6.0))
        
        if interpretaciones and 'carga_empleados' in interpretaciones:
            add_interpretation_docx(doc, interpretaciones['carga_empleados'])
            
    if 'timeline' in chart_paths:
        sec_timeline = doc.add_section()
        set_section_landscape(sec_timeline)
        
        h_t = doc.add_paragraph()
        run_ht = h_t.add_run("📈 Evolución Temporal del Volumen de Trabajo")
        run_ht.font.size = Pt(13)
        run_ht.font.bold = True
        run_ht.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        p_time = doc.add_paragraph()
        p_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_time.add_run().add_picture(chart_paths['timeline'], width=Inches(9.5))
        
        if interpretaciones and 'evolucion_temporal' in interpretaciones:
            add_interpretation_docx(doc, interpretaciones['evolucion_temporal'])
        
    sec_sla = doc.add_section()
    set_section_portrait(sec_sla)
    
    h_s = doc.add_paragraph()
    run_hs = h_s.add_run("⏱️ Métricas de Tiempos de Ciclo y SLA")
    run_hs.font.size = Pt(13)
    run_hs.font.bold = True
    run_hs.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    table_stats = doc.add_table(rows=6, cols=2)
    table_stats.allow_autofit = False
    table_stats.columns[0].width = Inches(3.75)
    table_stats.columns[1].width = Inches(3.75)
    for row in table_stats.rows:
        row.cells[0].width = Inches(3.75)
        row.cells[1].width = Inches(3.75)
    
    headers_stats = ["Métrica", "Valor (Horas)"]
    for j, text_h in enumerate(headers_stats):
        cell_h = table_stats.rows[0].cells[j]
        set_cell_background(cell_h, "E5E7EB")
        p_h = cell_h.paragraphs[0]
        r_h = p_h.add_run(text_h)
        r_h.font.bold = True
        r_h.font.size = Pt(9.5)
        
    stats_rows = [
        ("Tiempo Promedio", f"{ciclo_stats.get('promedio', 0):.2f} hrs"),
        ("Tiempo Mediana", f"{ciclo_stats.get('mediana', 0):.2f} hrs"),
        ("Tiempo Mínimo", f"{ciclo_stats.get('minima', 0):.2f} hrs"),
        ("Tiempo Máximo", f"{ciclo_stats.get('maxima', 0):.2f} hrs"),
        ("Cumplimiento SLA", f"{sla_cumplimiento:.2f}%")
    ]
    
    for idx_s, (name, val) in enumerate(stats_rows):
        row = table_stats.rows[idx_s + 1]
        c0 = row.cells[0]
        c0.paragraphs[0].text = ""
        c0.paragraphs[0].add_run(name).font.size = Pt(9.5)
        set_cell_margins(c0, top=40, bottom=40, left=80, right=80)
        
        c1 = row.cells[1]
        c1.paragraphs[0].text = ""
        c1.paragraphs[0].add_run(val).font.size = Pt(9.5)
        set_cell_margins(c1, top=40, bottom=40, left=80, right=80)
        
        bg_col = "FFFFFF" if idx_s % 2 == 0 else "F9FAFB"
        set_cell_background(c0, bg_col)
        set_cell_background(c1, bg_col)
        
    if interpretaciones and 'tiempos_ciclo' in interpretaciones:
        add_interpretation_docx(doc, interpretaciones['tiempos_ciclo'])
        
    if 'sla_histogram' in chart_paths:
        h_sh = doc.add_paragraph()
        run_hsh = h_sh.add_run("📊 Histograma de Distribución de Duraciones")
        run_hsh.font.size = Pt(11)
        run_hsh.font.bold = True
        run_hsh.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        p_hist = doc.add_paragraph()
        p_hist.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_hist.add_run().add_picture(chart_paths['sla_histogram'], width=Inches(6.0))
        
    if lista_variantes:
        sec_variants = doc.add_section()
        set_section_landscape(sec_variants)
        
        h_v = doc.add_paragraph()
        run_hv = h_v.add_run("🔄 Principales Variantes de Flujo Detectadas")
        run_hv.font.size = Pt(13)
        run_hv.font.bold = True
        run_hv.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        table_vars = doc.add_table(rows=1, cols=4)
        table_vars.allow_autofit = False
        table_vars.columns[0].width = Inches(1.0)
        table_vars.columns[1].width = Inches(1.0)
        table_vars.columns[2].width = Inches(1.0)
        table_vars.columns[3].width = Inches(7.0)
        
        headers_vars = ["Variante", "Casos", "Cobertura", "Secuencia de Actividades"]
        for j, text_h in enumerate(headers_vars):
            cell_h = table_vars.rows[0].cells[j]
            set_cell_background(cell_h, "E5E7EB")
            p_h = cell_h.paragraphs[0]
            r_h = p_h.add_run(text_h)
            r_h.font.bold = True
            r_h.font.size = Pt(9.5)
            
        for idx_v, v in enumerate(lista_variantes[:6]):
            row = table_vars.add_row()
            act_text = v.get('actividades')
            if isinstance(act_text, list):
                act_text = " ➔ ".join(act_text)
                
            vals_v = [
                v.get('id'),
                str(v.get('cantidad') or v.get('casos')),
                f"{v.get('porcentaje') or v.get('cobertura')}%",
                act_text
            ]
            for j, val_v in enumerate(vals_v):
                cell_v = row.cells[j]
                p_v = cell_v.paragraphs[0]
                p_v.add_run(val_v).font.size = Pt(9)
                set_cell_margins(cell_v, top=40, bottom=40, left=80, right=80)
                
                bg_col = "FFFFFF" if idx_v % 2 == 0 else "F9FAFB"
                set_cell_background(cell_v, bg_col)
                
        for row in table_vars.rows:
            row.cells[0].width = Inches(1.0)
            row.cells[1].width = Inches(1.0)
            row.cells[2].width = Inches(1.0)
            row.cells[3].width = Inches(7.0)
            
        if interpretaciones and 'variantes' in interpretaciones:
            add_interpretation_docx(doc, interpretaciones['variantes'])
                
    if matriz_traspaso and matriz_traspaso.get('values'):
        doc.add_paragraph()
        h_h = doc.add_paragraph()
        run_hh = h_h.add_run("🤝 Matriz de Traspaso (SNA)")
        run_hh.font.size = Pt(13)
        run_hh.font.bold = True
        run_hh.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        idx = matriz_traspaso.get('index', [])
        cols = matriz_traspaso.get('columns', [])
        vals = matriz_traspaso.get('values', [])
        
        original_len_cols = len(cols)
        original_len_idx = len(idx)
        if len(cols) > 8:
            cols = cols[:8]
        if len(idx) > 8:
            idx = idx[:8]
            
        table_sna = doc.add_table(rows=1, cols=len(cols) + 1)
        table_sna.allow_autofit = False
        
        num_cols = len(cols)
        col_width = Inches(8.4 / num_cols) if num_cols > 0 else Inches(1.0)
        
        cell_corner = table_sna.rows[0].cells[0]
        set_cell_background(cell_corner, "E5E7EB")
        cell_corner.paragraphs[0].add_run("De \\ A").font.bold = True
        cell_corner.paragraphs[0].runs[0].font.size = Pt(9)
        
        for j, c in enumerate(cols):
            cell_col = table_sna.rows[0].cells[j + 1]
            set_cell_background(cell_col, "E5E7EB")
            p_c = cell_col.paragraphs[0]
            p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_c = p_c.add_run(c)
            r_c.font.bold = True
            r_c.font.size = Pt(9)
            
        for i, row_name in enumerate(idx):
            row = table_sna.add_row()
            cell_row_lbl = row.cells[0]
            set_cell_background(cell_row_lbl, "F3F4F6")
            cell_row_lbl.paragraphs[0].text = ""
            cell_row_lbl.paragraphs[0].add_run(row_name).font.bold = True
            cell_row_lbl.paragraphs[0].runs[0].font.size = Pt(9)
            
            for j in range(len(cols)):
                val = vals[i][j]
                cell_val = row.cells[j + 1]
                p_val = cell_val.paragraphs[0]
                p_val.text = ""
                p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_val.add_run(str(int(val)) if val > 0 else "-").font.size = Pt(9)
                
        for row in table_sna.rows:
            row.cells[0].width = Inches(1.6)
            for j in range(len(cols)):
                row.cells[j + 1].width = col_width
                
        if original_len_cols > 8 or original_len_idx > 8:
            p_note = doc.add_paragraph()
            run_note = p_note.add_run("* Nota: La matriz se ha truncado a las primeras 8 columnas y filas para su visualización en el reporte.")
            run_note.font.size = Pt(7.5)
            run_note.font.italic = True
            run_note.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            
        if interpretaciones and 'matriz_traspaso' in interpretaciones:
            add_interpretation_docx(doc, interpretaciones['matriz_traspaso'])
                
    if slowest_cases:
        doc.add_paragraph()
        h_sc = doc.add_paragraph()
        run_hsc = h_sc.add_run("🚨 Detalle de Lotes Críticos (Mayor Duración)")
        run_hsc.font.size = Pt(13)
        run_hsc.font.bold = True
        run_hsc.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        table_sc = doc.add_table(rows=1, cols=5)
        table_sc.allow_autofit = False
        table_sc.columns[0].width = Inches(1.5)
        table_sc.columns[1].width = Inches(2.2)
        table_sc.columns[2].width = Inches(2.2)
        table_sc.columns[3].width = Inches(1.3)
        table_sc.columns[4].width = Inches(2.8)
        
        headers_sc = ["ID Lote (Caso)", "Inicio", "Fin", "Duración", "Estado SLA"]
        for j, text_h in enumerate(headers_sc):
            cell_h = table_sc.rows[0].cells[j]
            set_cell_background(cell_h, "E5E7EB")
            p_h = cell_h.paragraphs[0]
            r_h = p_h.add_run(text_h)
            r_h.font.bold = True
            r_h.font.size = Pt(9.5)
            
        for idx_sc, c in enumerate(slowest_cases[:5]):
            row = table_sc.add_row()
            vals_sc = [
                c.get('id_caso'),
                c.get('inicio'),
                c.get('fin'),
                c.get('duracion'),
                c.get('status')
            ]
            for j, val_sc in enumerate(vals_sc):
                cell_v = row.cells[j]
                p_v = cell_v.paragraphs[0]
                run_v = p_v.add_run(val_sc)
                run_v.font.size = Pt(9)
                if j == 4 and "retraso" in val_sc.lower():
                    run_v.font.bold = True
                    run_v.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                set_cell_margins(cell_v, top=40, bottom=40, left=80, right=80)
                
                bg_col = "FFFFFF" if idx_sc % 2 == 0 else "F9FAFB"
                set_cell_background(cell_v, bg_col)
                
        for row in table_sc.rows:
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(2.2)
            row.cells[2].width = Inches(2.2)
            row.cells[3].width = Inches(1.3)
            row.cells[4].width = Inches(2.8)
            
        if interpretaciones and 'lotes_criticos' in interpretaciones:
            add_interpretation_docx(doc, interpretaciones['lotes_criticos'])
                
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


